from docx import Document
from io import BytesIO
from random import randint

class CitationFixer:
    def __init__(self):
        pass
    def read_document(self, path):
        with open(path, 'rb') as f:
            source_stream = BytesIO(f.read())
        self.document = Document(source_stream)
        self.citations = [[] for _ in range(len(self.document.paragraphs))]
    def get_citations(self):
        for i in range(len(self.document.paragraphs)):
            text = self.document.paragraphs[i].text
            if len(text) > 6:
                start_char = randint(0, len(text) - 6)
                end_char = start_char + 5
                self.citations[i].append((start_char, end_char))
    def fix_citations(self):
        for i in range(len(self.document.paragraphs)):
            start_char, end_char = self.citations[i]
            text = self.document.paragraphs[i].text
            self.document.paragraphs[i].text = text[:start_char] + 'a' * (end_char-start_char+1) + text[end_char+1:]
    def save_document(self, path):
        self.document.save(path)

class SimpleNLPCitationFixer(CitationFixer):
    def __init__(self):
        self.min_paragraph_length = 10
        self.nlp = spacy.load("en_core_web_sm")
    def get_citations(self):
        print("inside get_citations nlp")
        print("len[self.citations]:", len(self.citations))
        for i in range(len(self.document.paragraphs)):
            text = self.document.paragraphs[i].text
            if len(text) >= self.min_paragraph_length:
                entities = self.nlp(text).ents
                for j in range(len(entities)-1):
                    if entities[j].label_ == 'PERSON' and entities[j+1].label_ == 'DATE':
                        start_char = entities[j].start_char
                        end_char = entities[j+1].end_char
                        # spacy gives the intervals in the [start_char, end_char) format,
                        # but we want it to be in the [start_char, end_char] format.
                        # Since we want to include ')', we don't decrement end_char
                        if text[end_char-len(entities[j+1].text)-1] == '(':
                            # ensuring that there is a ( before the year (distinguishing it from cases
                            # there is a year but not for the purpose of a citation)
                            self.citations[i].append((start_char, end_char))
#                             print("last character:")
#                             print("citation:", text[start_char:end_char+1])